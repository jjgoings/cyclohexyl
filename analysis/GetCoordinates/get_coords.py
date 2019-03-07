import cclib
import pandas as pd
import docx
#import seaborn as sns
import numpy as np
#import matplotlib.pyplot as plt
import os
from docx.shared import Pt

#sns.set_color_codes()

class Structure(object):
    def __init__(self,logfile):
        self.data     = cclib.io.ccread(logfile)
        self.energy   = self.data.scfenergies[-1]/27.21138505
        self.coords   = self.data.atomcoords[-1]
        self.atoms    = []
        for atom in self.data.atomnos:
            self.atoms.append(self.num2sym(atom))
        assert len(self.atoms) == len(self.coords)

    def num2sym(self,atomno):
        """Routine that converts atomic number to atomic symbol"""
        symbol = [
            "X","H","He",
            "Li","Be","B","C","N","O","F","Ne",
            "Na","Mg","Al","Si","P","S","Cl","Ar",
            "K", "Ca", "Sc", "Ti", "V", "Cr", "Mn", "Fe",
            "Co", "Ni", "Cu", "Zn",
            "Ga", "Ge", "As", "Se", "Br", "Kr",
            "Rb", "Sr", "Y", "Zr", "Nb", "Mo", "Tc", "Ru",
            "Rh", "Pd", "Ag", "Cd",
            "In", "Sn", "Sb", "Te", "I", "Xe",
            "Cs", "Ba", "La", "Ce", "Pr", "Nd", "Pm", "Sm",  "Eu",
            "Gd", "Tb", "Dy", "Ho", "Er", "Tm", "Yb", "Lu",
            "Hf", "Ta", "W", "Re", "Os", "Ir", "Pt", "Au", "Hg",
            "Tl","Pb","Bi","Po","At","Rn"]
        return symbol[atomno]
        
    
class Molecule(object):
    def __init__(self,prefix):
        self.prefix  = str(prefix)
        path = '../../'+self.prefix+'BIP/'
        self.states = []

        for state in ['neutral','E0PT','E1PT','E2PT','E3PT','E4PT']:
            if prefix in ['amide-mono','unsub-mono']:
                logfile = path+state+'/opt_freq/'+self.prefix+'BIP-'+state+'.log'
            else:
                logfile = path+state+'/opt_freq/cyclohexyl-'+self.prefix+'BIP-'+state+'.log'
            if os.path.isfile(logfile):
                setattr(self,state.lower(),Structure(logfile))
                setattr(getattr(self,state.lower()),'name',self.prefix+'BIP-cyclohexylimine'+': '+state+'.')
                self.states.append(state)
               
    def get_structures(self):
        #print('Checking ',self.prefix+'BIP-cyclohexylimine structures...')
        doc = docx.Document('./template.docx') # read in tempalte document with table style defined
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        for state,data in [(x,getattr(self,x.lower())) for x in self.states]: 
            df = pd.DataFrame(data={'atom': data.atoms, 'x':data.coords[:,0],'y':data.coords[:,1],'z':data.coords[:,2]})
            #doc.add_paragraph('Table S#. Cartesian coordinates for %s (%.8f Hartrees)' % (data.name,data.energy))
            p = doc.add_paragraph()
            runner = p.add_run('Table S#.') 
            runner.bold = True
            runner2 = p.add_run(' Cartesian coordinates for %s (%.8f Hartrees)' % (data.name,data.energy))
            t = doc.add_table(df.shape[0]+1, df.shape[1])
            t.style = 'jjg2' # must be defined in a template document...here called "template.docx"
            
            # add the header rows.
            for j in range(df.shape[-1]):
                t.cell(0,j).text = df.columns[j]
            
            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    if j == 0:
                        t.cell(i+1,j).text = str(df.values[i,j])
                    else:
                        t.cell(i+1,j).text = "{:10.6f}".format(df.values[i,j])
            
            doc.add_page_break() 
            # save the doc
        doc.save(self.prefix+'BIP_structures.docx')



#            with open(filename,'w') as f:
#                f.write('Cartesian coordinates for %s (%.8f Hartrees)\n' % (data.name,data.energy))
#                f.write('%4s, %12s, %12s, %12s\n' % ('atom','x','y','z'))
#                for i in range(len(data.atoms)):
#                    f.write('%4s, %12.6f, %12.6f, %12.6f\n' % (data.atoms[i],data.coords[i,0],data.coords[i,1],data.coords[i,2]))
         
            

if __name__ == '__main__':
    amidemono = Molecule('amide-mono')
    amidemono.get_structures()
    unsubmono = Molecule('unsub-mono')
    unsubmono.get_structures()
    mono = Molecule('mono')
    mono.get_structures()
    di = Molecule('di')
    di.get_structures() 
    tri = Molecule('tri')
    tri.get_structures()
    


